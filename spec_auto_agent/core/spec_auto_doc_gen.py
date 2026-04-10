"""
spec_auto_doc_gen.py
─────────────────────────────────────────────────────────
SpecAutoAgent - Step 2: 규격서 문서 생성 (Excel / Word)
SpecApiSchema → API 연동 규격서 파일 자동 생성
─────────────────────────────────────────────────────────
"""

import os
from datetime import datetime
from pathlib import Path

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from spec_auto_agent.models.spec_auto_models import SpecApiSchema

OUTPUT_DIR = Path(__file__).parent.parent / "output"

COLOR_HEADER_BG = "1F4E79"
COLOR_HEADER_FG = "FFFFFF"
COLOR_REQUIRED  = "FF0000"
COLOR_OPTIONAL  = "888888"


class SpecAutoDocGen:
    """
    SpecApiSchema → Excel / Word 규격서 자동 생성
    SpecAutoAgent Step 2 문서화 담당 컴포넌트
    """

    def __init__(self, author: str = "SpecAutoAgent"):
        self.author = author
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # ── Public ──────────────────────────────────

    def generate_excel(self, schema: SpecApiSchema) -> str:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "API 규격서"

        for i, w in enumerate([5, 20, 15, 10, 15, 35, 20], 1):
            ws.column_dimensions[get_column_letter(i)].width = w

        row = self._title(ws, schema, 1) + 1
        row = self._basic_info(ws, schema, row) + 1
        row = self._section(ws, "📥 Request Parameters", row)
        row = self._field_table(ws, schema.request, row) + 1
        row = self._section(ws, "📤 Response Parameters", row)
        row = self._field_table(ws, schema.response, row) + 1
        row = self._section(ws, "⚠️ Error Codes", row)
        self._error_table(ws, schema.errors, row)

        path = str(OUTPUT_DIR / f"spec_auto_{schema.api_name.replace(' ', '_')}_{self._ts()}.xlsx")
        wb.save(path)
        print(f"[SpecAutoDocGen] Excel 생성 완료: {path}")
        return path

    def generate_word(self, schema: SpecApiSchema) -> str:
        doc = Document()
        for sec in doc.sections:
            sec.top_margin = sec.bottom_margin = Cm(2)
            sec.left_margin = sec.right_margin = Cm(2.5)

        t = doc.add_heading("API 연동 규격서", level=0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub = doc.add_paragraph(f"{schema.api_name}  |  작성자: {self.author}  |  {self._ts()}")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        doc.add_heading("1. 기본 정보", level=1)
        tbl = doc.add_table(rows=5, cols=2)
        tbl.style = "Table Grid"
        for i, (k, v) in enumerate([
            ("API 명칭",    schema.api_name),
            ("엔드포인트", f"{schema.method}  {schema.endpoint}"),
            ("설명",        schema.description),
            ("버전",        schema.version),
            ("인증 필요",   "Y" if schema.auth_required else "N"),
        ]):
            tbl.rows[i].cells[0].text = k
            tbl.rows[i].cells[1].text = str(v)
            tbl.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        doc.add_paragraph()

        doc.add_heading("2. Request Parameters", level=1)
        self._word_field_table(doc, schema.request)
        doc.add_paragraph()

        doc.add_heading("3. Response Parameters", level=1)
        self._word_field_table(doc, schema.response)
        doc.add_paragraph()

        doc.add_heading("4. Error Codes", level=1)
        et = doc.add_table(rows=1 + len(schema.errors), cols=2)
        et.style = "Table Grid"
        et.rows[0].cells[0].text = "Error Code"
        et.rows[0].cells[1].text = "설명"
        for run in [et.rows[0].cells[0].paragraphs[0].runs[0],
                    et.rows[0].cells[1].paragraphs[0].runs[0]]:
            run.bold = True
        for i, e in enumerate(schema.errors, 1):
            et.rows[i].cells[0].text = e.code
            et.rows[i].cells[1].text = e.description

        path = str(OUTPUT_DIR / f"spec_auto_{schema.api_name.replace(' ', '_')}_{self._ts()}.docx")
        doc.save(path)
        print(f"[SpecAutoDocGen] Word 생성 완료: {path}")
        return path

    # ── Private: Excel helpers ───────────────────

    def _title(self, ws, schema, row):
        ws.merge_cells(f"A{row}:G{row}")
        c = ws[f"A{row}"]
        c.value = f"API 연동 규격서 — {schema.api_name}"
        c.font = Font(name="맑은 고딕", bold=True, size=16, color=COLOR_HEADER_FG)
        c.fill = PatternFill("solid", fgColor=COLOR_HEADER_BG)
        c.alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[row].height = 30
        return row + 1

    def _basic_info(self, ws, schema, row):
        for label, val in [
            ("엔드포인트", f"{schema.method}  {schema.endpoint}"),
            ("설명",        schema.description),
            ("버전",        schema.version),
            ("작성자",      self.author),
            ("작성일",      self._ts()),
            ("인증",        "필수" if schema.auth_required else "불필요"),
        ]:
            ws[f"B{row}"] = label
            ws[f"B{row}"].font = Font(bold=True, name="맑은 고딕")
            ws[f"C{row}"] = val
            row += 1
        return row

    def _section(self, ws, title, row):
        ws.merge_cells(f"A{row}:G{row}")
        c = ws[f"A{row}"]
        c.value = title
        c.font = Font(bold=True, size=12, color="FFFFFF", name="맑은 고딕")
        c.fill = PatternFill("solid", fgColor="2E75B6")
        c.alignment = Alignment(horizontal="left", vertical="center", indent=1)
        ws.row_dimensions[row].height = 22
        return row + 1

    def _field_table(self, ws, fields, row):
        thin = Side(style="thin", color="AAAAAA")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)
        for col, h in enumerate(["#", "필드명", "타입", "필수", "예시", "설명", "포맷"], 1):
            c = ws.cell(row=row, column=col, value=h)
            c.font = Font(bold=True, color="FFFFFF", name="맑은 고딕")
            c.fill = PatternFill("solid", fgColor="4472C4")
            c.alignment = Alignment(horizontal="center")
        row += 1
        for i, f in enumerate(fields, 1):
            for col, val in enumerate([i, f.name, f.type,
                                        "Y" if f.required else "N",
                                        str(f.example or ""), f.description,
                                        f.format or ""], 1):
                c = ws.cell(row=row, column=col, value=val)
                c.border = border
                c.alignment = Alignment(horizontal="center" if col in [1, 4] else "left",
                                        vertical="center")
                if col == 4:
                    c.font = Font(color=COLOR_REQUIRED if val == "Y" else COLOR_OPTIONAL,
                                  bold=True, name="맑은 고딕")
            row += 1
        return row

    def _error_table(self, ws, errors, row):
        for col, h in enumerate(["#", "Error Code", "설명"], 1):
            c = ws.cell(row=row, column=col, value=h)
            c.font = Font(bold=True, color="FFFFFF", name="맑은 고딕")
            c.fill = PatternFill("solid", fgColor="C00000")
            c.alignment = Alignment(horizontal="center")
        row += 1
        for i, e in enumerate(errors, 1):
            ws.cell(row=row, column=1, value=i)
            ws.cell(row=row, column=2, value=e.code)
            ws.cell(row=row, column=3, value=e.description)
            row += 1
        return row

    # ── Private: Word helpers ───────────────────

    def _word_field_table(self, doc, fields):
        if not fields:
            doc.add_paragraph("(없음)")
            return
        headers = ["필드명", "타입", "필수", "예시", "설명"]
        tbl = doc.add_table(rows=1 + len(fields), cols=len(headers))
        tbl.style = "Table Grid"
        for col, h in enumerate(headers):
            tbl.rows[0].cells[col].text = h
            tbl.rows[0].cells[col].paragraphs[0].runs[0].bold = True
        for i, f in enumerate(fields, 1):
            tbl.rows[i].cells[0].text = f.name
            tbl.rows[i].cells[1].text = f.type
            tbl.rows[i].cells[2].text = "Y" if f.required else "N"
            tbl.rows[i].cells[3].text = str(f.example or "")
            tbl.rows[i].cells[4].text = f.description

    def _ts(self):
        return datetime.now().strftime("%Y%m%d_%H%M%S")
