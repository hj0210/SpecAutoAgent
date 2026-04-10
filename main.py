"""
main.py
─────────────────────────────────────────────────────────
SpecAutoAgent - FastAPI 진입점

URL 구조:
  /          → /ui 로 리다이렉트
  /ui        → 프롬프트 입력 Web UI
  /docs      → Swagger API 문서 (자동 생성)
  /health    → 헬스체크

  /api/spec-auto/generate            → 자연어 단일 입력
  /api/spec-auto/generate-structured → 구조화 입력 (UI 전용)
  /api/spec-auto/analyze             → 스펙 분석만 (파일 생성 X)
  /api/spec-auto/download/{filename} → 생성 파일 다운로드
  /api/spec-auto/parse-reference     → 기존 규격서 파일 → 텍스트 추출
─────────────────────────────────────────────────────────
"""

import io
import os
from pathlib import Path

from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, RedirectResponse, HTMLResponse
from fastapi.templating import Jinja2Templates
from fastapi.requests import Request
from dotenv import load_dotenv

import openpyxl
from docx import Document as DocxDocument

from spec_auto_agent.models.spec_auto_models import (
    SpecAutoRequest,
    SpecAutoStructuredRequest,
    SpecAutoJavaRequest,
    SpecAutoResponse,
    SpecAutoAssets,
    SpecAutoHealthResponse,
    OutputFormat,
)
from spec_auto_agent.core.spec_auto_analyzer import SpecAutoAnalyzer
from spec_auto_agent.core.spec_auto_doc_gen  import SpecAutoDocGen
from spec_auto_agent.core.spec_auto_code_gen import SpecAutoCodeGen
from spec_auto_agent.core.spec_auto_postman  import SpecAutoPostman

load_dotenv()

OUTPUT_DIR  = Path(__file__).parent / "spec_auto_agent" / "output"
TEMPLATE_DIR = Path(__file__).parent / "templates"

# ─────────────────────────────────────────────
# FastAPI 앱
# ─────────────────────────────────────────────

app = FastAPI(
    title       = "SpecAutoAgent",
    description = "Natural Language → API Spec, Docs & Code. Automated.",
    version     = "1.0.0",
    docs_url    = "/docs",
    redoc_url   = "/redoc",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], allow_credentials=True,
    allow_methods=["*"], allow_headers=["*"],
)

templates = Jinja2Templates(directory=str(TEMPLATE_DIR))

# 서비스 인스턴스
analyzer = SpecAutoAnalyzer()
postman  = SpecAutoPostman()


# ─────────────────────────────────────────────
# 페이지 라우터
# ─────────────────────────────────────────────

@app.get("/", include_in_schema=False)
def root():
    """루트 → UI 페이지로 리다이렉트"""
    return RedirectResponse(url="/ui")


@app.get("/ui", response_class=HTMLResponse, include_in_schema=False)
def ui_page(request: Request):
    """프롬프트 입력 Web UI"""
    return templates.TemplateResponse("index.html", {"request": request})


@app.get("/health", response_model=SpecAutoHealthResponse, tags=["Health"])
def health():
    return SpecAutoHealthResponse()


# ─────────────────────────────────────────────
# API 라우터
# ─────────────────────────────────────────────

@app.post("/api/spec-auto/generate", response_model=SpecAutoResponse, tags=["SpecAutoAgent"])
async def spec_auto_generate(request: SpecAutoRequest):
    """자연어 단일 텍스트 입력 → 규격서 & 코드 생성 (Swagger / curl 용)"""
    try:
        schema = analyzer.analyze(request.user_input)
    except RuntimeError as e:
        raise HTTPException(status_code=400, detail=str(e))
    return await _run_pipeline(schema=schema, request=request)


@app.post("/api/spec-auto/generate-structured", response_model=SpecAutoResponse, tags=["SpecAutoAgent"])
async def spec_auto_generate_structured(request: SpecAutoStructuredRequest):
    """
    ## 구조화 입력 → 규격서 & 코드 생성 (Web UI 전용)

    목적 / 기능 / 필드 / 특이사항을 분리하여 입력합니다.
    기존 규격서에서 추출한 `reference` 텍스트를 넣으면 참고하여 스펙을 설계합니다.
    """
    try:
        schema = analyzer.analyze_structured(
            purpose   = request.purpose,
            features  = request.features,
            fields    = request.fields,
            notes     = request.notes,
            reference = request.reference,
        )
    except RuntimeError as e:
        raise HTTPException(status_code=400, detail=str(e))
    return await _run_pipeline(schema=schema, request=request)


@app.post("/api/spec-auto/parse-java", response_model=SpecAutoResponse, tags=["SpecAutoAgent"])
async def spec_auto_parse_java(request: SpecAutoJavaRequest):
    """
    ## Java Spring Controller 코드 → 규격서 & 코드 생성 (Scenario A)

    업로드된 `.java` 파일의 내용을 직접 전달하면 Regex + GPT-4o 하이브리드로 분석합니다.
    기존 규격서 텍스트(`reference_doc`)를 함께 제공하면 패턴을 맞춰 규격서를 생성합니다.
    """
    try:
        schema = analyzer.analyze_from_java(
            java_code     = request.java_code,
            reference_doc = request.reference_doc,
        )
    except RuntimeError as e:
        raise HTTPException(status_code=400, detail=str(e))
    return await _run_pipeline(schema=schema, request=request)


@app.post("/api/spec-auto/parse-java-file", response_model=SpecAutoResponse, tags=["SpecAutoAgent"])
async def spec_auto_parse_java_file(
    file:          UploadFile = File(...),
    reference_doc: str        = "",
    author:        str        = "SpecAutoAgent",
):
    """
    ## Java 파일 직접 업로드 → 규격서 & 코드 생성 (Scenario A - UI용)

    `.java` 파일을 multipart/form-data로 업로드합니다.
    """
    if not (file.filename or "").endswith(".java"):
        raise HTTPException(status_code=400, detail=".java 파일만 지원합니다.")

    content = await file.read()
    try:
        java_code = content.decode("utf-8")
    except UnicodeDecodeError:
        java_code = content.decode("euc-kr", errors="replace")

    schema = analyzer.analyze_from_java(java_code=java_code, reference_doc=reference_doc)

    # _run_pipeline은 output_formats / code_language / author 속성을 사용하므로 임시 객체 생성
    class _Req:
        output_formats = [OutputFormat.ALL]
        code_language  = "java"
        purpose        = schema.api_name

    _req = _Req()
    _req.author = author
    return await _run_pipeline(schema=schema, request=_req)


@app.post("/api/spec-auto/analyze", tags=["SpecAutoAgent"])
async def spec_auto_analyze(user_input: str):
    """자연어 → JSON Schema 분석 결과만 반환 (파일 생성 없음)"""
    try:
        schema = analyzer.analyze(user_input)
        return {"success": True, "api_schema": schema.model_dump()}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/spec-auto/download/{filename}", tags=["SpecAutoAgent"])
def spec_auto_download(filename: str):
    """생성된 산출물 파일 다운로드"""
    filepath = OUTPUT_DIR / filename
    if not filepath.exists():
        raise HTTPException(status_code=404, detail=f"파일을 찾을 수 없습니다: {filename}")
    # 경로 탈출 방지
    if not filepath.resolve().is_relative_to(OUTPUT_DIR.resolve()):
        raise HTTPException(status_code=403, detail="접근 거부")
    return FileResponse(
        path=str(filepath),
        filename=filename,
        media_type="application/octet-stream",
    )


@app.post("/api/spec-auto/parse-reference", tags=["SpecAutoAgent"])
async def parse_reference(file: UploadFile = File(...)):
    """
    ## 기존 규격서 파일 업로드 → 텍스트 추출

    Excel(.xlsx) 또는 Word(.docx) 규격서를 업로드하면
    텍스트를 추출하여 반환합니다.
    이 텍스트를 generate-structured의 `reference` 필드에 넣어 참고 문서로 활용하세요.
    """
    filename = file.filename or ""
    content  = await file.read()

    try:
        if filename.endswith(".xlsx"):
            text = _extract_excel_text(content)
        elif filename.endswith(".docx"):
            text = _extract_word_text(content)
        else:
            raise HTTPException(status_code=400, detail=".xlsx 또는 .docx 파일만 지원합니다.")

        return {
            "success":  True,
            "filename": filename,
            "text":     text,
            "length":   len(text),
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"파일 파싱 오류: {e}")


# ─────────────────────────────────────────────
# 내부 헬퍼
# ─────────────────────────────────────────────

async def _run_pipeline(schema, request) -> SpecAutoResponse:
    """공통 파이프라인: schema → 산출물 생성"""
    try:
        print(f"[SpecAutoAgent] 파이프라인 시작 → {schema.method} {schema.endpoint}")
        assets  = SpecAutoAssets()
        all_fmt = OutputFormat.ALL in request.output_formats
        doc_gen = SpecAutoDocGen(author=request.author or "SpecAutoAgent")

        if all_fmt or OutputFormat.EXCEL in request.output_formats:
            assets.excel_path = Path(doc_gen.generate_excel(schema)).name

        if all_fmt or OutputFormat.WORD in request.output_formats:
            assets.word_path = Path(doc_gen.generate_word(schema)).name

        if all_fmt or OutputFormat.POSTMAN in request.output_formats:
            pdict, ppath        = postman.build(schema)
            assets.postman_path = Path(ppath).name
            assets.postman_json = pdict

        if all_fmt or OutputFormat.CODE in request.output_formats:
            code, cpath             = SpecAutoCodeGen().generate(schema, request.code_language)
            assets.sample_code_path = Path(cpath).name
            assets.sample_code      = code

        print(f"[SpecAutoAgent] 완료 ✅")
        return SpecAutoResponse(
            success    = True,
            user_input = getattr(request, "user_input", request.purpose),
            api_schema = schema,
            assets     = assets,
            message    = f"'{schema.api_name}' 규격서 및 코드 생성 완료",
        )
    except Exception as e:
        print(f"[SpecAutoAgent] 오류: {e}")
        raise HTTPException(status_code=500, detail=str(e))


def _extract_excel_text(content: bytes) -> str:
    """Excel 파일에서 셀 텍스트 추출"""
    wb    = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
    lines = []
    for ws in wb.worksheets:
        lines.append(f"[시트: {ws.title}]")
        for row in ws.iter_rows(values_only=True):
            row_text = "\t".join(str(c) if c is not None else "" for c in row)
            if row_text.strip():
                lines.append(row_text)
    return "\n".join(lines)


def _extract_word_text(content: bytes) -> str:
    """Word 파일에서 단락 + 표 텍스트 추출"""
    doc   = DocxDocument(io.BytesIO(content))
    lines = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                lines.append(row_text)
    return "\n".join(lines)
